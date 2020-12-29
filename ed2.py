import docx 
from docx.shared import Pt
import pandas as pd

file = input("Enter path to file or 'filename.docx' : ")
doc = docx.Document(file) 
print("Amount of paragraphs: ", len(doc.paragraphs))
par1 = doc.paragraphs[0]
t = []
sizes = []
names = []
bold = []
color = []
italic = []
count = 1

for par1 in doc.paragraphs:
  for run in par1.runs:
    t.append(run.text)

    if run.font.size is not None:
     sizes.append(str(run.font.size))

    if run.font.name is not None:
     names.append(str(run.font.name))

    if run.font.bold is not None:
     bold.append(str(run.font.bold))

    if run.font.color.rgb is not None:
     color.append(str(run.font.color.rgb))

    if run.font.italic is not None:
     italic.append(str(run.font.italic))

  print(' '.join(t))
  print("\nParagraph",count,":\n===========")
  if len(sizes) == 0:
   print("No unique font sizes, except default")
  else:
   print("Sizes are:")
   print(' '.join(pd.unique(pd.Series(sizes))))

  if len(names) == 0:
   print("No unique font names, except default")
  else:
   print("Font names are:")
   print(' '.join(pd.unique(pd.Series(names))))

  if len(bold) == 0:
   print("No unique bold text")
  else:
   print("Text include bold font:")
   print(' '.join(pd.unique(pd.Series(bold))))

  if len(color) == 0:
   print("No unique font colors, except default")
  else:
   print("Colors are:")
   print(' '.join(pd.unique(pd.Series(color))))

  if len(italic) == 0:
   print("No unique italic text")
  else:
   print("Text include italic font:")
   print(' '.join(pd.unique(pd.Series(italic))))
  
  count+=1

