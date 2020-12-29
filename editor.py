# 
import docx 
from docx.shared import Pt
from docx.shared import RGBColor

doc = docx.Document()



run = doc.add_paragraph().add_run("Оказавшись в приюте, 9-летняя Бет демонстрирует поразительный талант к "
                                       "шахматам и сталкивается с растущей зависимостью от выдаваемых детям"
                                       " транквилизаторов.")

style = doc.styles['Normal']
fontS = style.font
fontS.name = 'Bookman Old Style'
fontS.size = Pt(20)
fontS.italic = True
fontS.color.rgb = RGBColor(0x42, 0x24, 0xE9)

doc.save('f.docx')